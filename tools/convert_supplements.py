#!/usr/bin/env python3
"""
Convert project markdown supplements to:
  1. Styled DOCX with theme-driven styling + full-page cover image
  2. Triple-spaced PDF for editing/annotation

Theme is loaded from styles/layout/<layout.docx_theme>.theme.json, selected
via config/system.json (see load_theme() below).

Requires: python-docx, Pillow, pandoc, weasyprint
"""

import json
import os
import re
import subprocess
import tempfile
from pathlib import Path

from docx import Document
from docx.shared import Emu, Pt, RGBColor, Twips
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from PIL import Image

REPO_ROOT = Path(__file__).resolve().parent.parent
PROJECTS_DIR = REPO_ROOT / "projects"

# 1 twip = 1/20 pt = 1/1440 inch = 635 EMU
EMU_PER_TWIP = 635


def load_theme():
    """Load the DOCX theme selected by config/system.json -> layout.docx_theme.

    Mirrors the loader in scripts/export-docx.js: config -> theme name ->
    styles/layout/<name>.theme.json. Falls back to the "default" theme name
    if config/system.json is missing or doesn't specify one.
    """
    config_path = REPO_ROOT / "config" / "system.json"
    try:
        cfg = json.loads(config_path.read_text(encoding="utf-8"))
    except (OSError, json.JSONDecodeError):
        cfg = {}

    theme_name = (cfg.get("layout") or {}).get("docx_theme") or "default"
    theme_path = REPO_ROOT / "styles" / "layout" / f"{theme_name}.theme.json"
    theme = json.loads(theme_path.read_text(encoding="utf-8"))

    publisher_line = (cfg.get("system") or {}).get("publisher_line") or ""
    return theme, publisher_line


THEME, PUBLISHER_LINE = load_theme()


def _rgb(hex_str):
    return RGBColor.from_string(hex_str)


# Colors/fonts pulled from the active theme. Mirrors scripts/export-docx.js's
# heading convention: Title/Heading 1/Heading 2 use accent1, Subtitle/
# Heading 3/Heading 4/Heading 5 use accent2 — so this reference template
# renders headings in the same colors as the JS DOCX exporter for the same
# theme.
BODY_COLOR = _rgb(THEME["colors"]["body"])
ACCENT1_COLOR = _rgb(THEME["colors"]["accent1"])
ACCENT2_COLOR = _rgb(THEME["colors"]["accent2"])

BODY_FONT = THEME["fonts"]["body"]
HEADING_FONT = THEME["fonts"]["heading"]

# CSS for triple-spaced PDF styling
PDF_CSS = """
@page {
    size: 8.5in 11in;
    margin: 1in 1in 1in 1in;
    @bottom-center { content: counter(page); font-family: 'P052', serif; font-size: 9pt; color: #666; }
}

@page :first { margin: 0; @bottom-center { content: none; } }

body {
    font-family: 'P052', 'Palatino', 'Palatino Linotype', 'URW Palladio L', serif;
    font-size: 11pt;
    color: #1a1a1a;
    line-height: 3.0;
}

h1, h2, h3, h4, h5, h6 { line-height: 1.3; }

h1 {
    font-family: 'URW Bookman', 'Bookman Old Style', serif;
    font-size: 22pt;
    font-weight: bold;
    color: #8B0000;
    page-break-before: always;
    margin-top: 24pt;
    margin-bottom: 12pt;
}

h1:first-of-type { page-break-before: avoid; }

h2 {
    font-family: 'URW Bookman', 'Bookman Old Style', serif;
    font-size: 16pt;
    font-weight: bold;
    color: #4A002A;
    margin-top: 18pt;
    margin-bottom: 8pt;
}

h3 {
    font-family: 'URW Bookman', 'Bookman Old Style', serif;
    font-size: 13pt;
    font-weight: bold;
    color: #2F1E0E;
    margin-top: 12pt;
    margin-bottom: 6pt;
}

h4 {
    font-family: 'URW Bookman', 'Bookman Old Style', serif;
    font-size: 12pt;
    font-weight: bold;
    font-style: italic;
    color: #2F1E0E;
}

blockquote {
    font-style: italic;
    color: #444;
    border-left: 3px solid #8B0000;
    padding-left: 1em;
    margin-left: 0;
}

table {
    border-collapse: collapse;
    width: 100%;
    margin: 12pt 0;
    line-height: 1.4;
    font-size: 10pt;
}

th, td {
    border: 1px solid #999;
    padding: 4pt 8pt;
    text-align: left;
}

th {
    background-color: #4A002A;
    color: white;
    font-weight: bold;
}

img {
    max-width: 100%;
    height: auto;
    display: block;
    margin: 12pt auto;
}

hr {
    border: none;
    border-top: 2px solid #8B0000;
    margin: 24pt 0;
}

code {
    font-family: 'DejaVu Sans Mono', monospace;
    font-size: 9pt;
    background-color: #f5f5f5;
    padding: 1pt 3pt;
}

pre {
    font-family: 'DejaVu Sans Mono', monospace;
    font-size: 9pt;
    background-color: #f5f5f5;
    padding: 8pt;
    line-height: 1.4;
    white-space: pre-wrap;
    word-wrap: break-word;
}

.cover-image {
    page: first;
    width: 8.5in;
    height: 11in;
    object-fit: cover;
    page-break-after: always;
}
"""


def create_reference_docx(output_path):
    """Create a pandoc reference.docx with theme-driven styling."""
    doc = Document()

    section = doc.sections[0]
    section.page_width = Twips(THEME["page"]["w"])
    section.page_height = Twips(THEME["page"]["h"])
    section.left_margin = Twips(THEME["page"]["ml"])
    section.right_margin = Twips(THEME["page"]["mr"])
    section.top_margin = Twips(THEME["page"]["mt"])
    section.bottom_margin = Twips(THEME["page"]["mb"])

    # Normal
    style = doc.styles['Normal']
    style.font.name = BODY_FONT
    style.font.size = Pt(11)
    style.font.color.rgb = BODY_COLOR
    pf = style.paragraph_format
    pf.space_after = Pt(6)
    pf.space_before = Pt(0)
    pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    pf.line_spacing = 3.0

    # Title
    style = doc.styles['Title']
    style.font.name = HEADING_FONT
    style.font.size = Pt(28)
    style.font.bold = True
    style.font.color.rgb = ACCENT1_COLOR
    style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    style.paragraph_format.space_after = Pt(12)
    style.paragraph_format.space_before = Pt(24)

    # Subtitle
    style = doc.styles['Subtitle']
    style.font.name = BODY_FONT
    style.font.size = Pt(14)
    style.font.italic = True
    style.font.color.rgb = ACCENT2_COLOR
    style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    style.paragraph_format.space_after = Pt(12)

    # Heading 1
    style = doc.styles['Heading 1']
    style.font.name = HEADING_FONT
    style.font.size = Pt(22)
    style.font.bold = True
    style.font.color.rgb = ACCENT1_COLOR
    style.paragraph_format.space_before = Pt(24)
    style.paragraph_format.space_after = Pt(12)
    style.paragraph_format.page_break_before = True

    # Heading 2
    style = doc.styles['Heading 2']
    style.font.name = HEADING_FONT
    style.font.size = Pt(16)
    style.font.bold = True
    style.font.color.rgb = ACCENT1_COLOR
    style.paragraph_format.space_before = Pt(18)
    style.paragraph_format.space_after = Pt(8)

    # Heading 3
    style = doc.styles['Heading 3']
    style.font.name = HEADING_FONT
    style.font.size = Pt(13)
    style.font.bold = True
    style.font.color.rgb = ACCENT2_COLOR
    style.paragraph_format.space_before = Pt(12)
    style.paragraph_format.space_after = Pt(6)

    # Heading 4
    style = doc.styles['Heading 4']
    style.font.name = HEADING_FONT
    style.font.size = Pt(12)
    style.font.bold = True
    style.font.italic = True
    style.font.color.rgb = ACCENT2_COLOR
    style.paragraph_format.space_before = Pt(10)
    style.paragraph_format.space_after = Pt(4)

    # Heading 5
    style = doc.styles['Heading 5']
    style.font.name = HEADING_FONT
    style.font.size = Pt(11)
    style.font.bold = True
    style.font.color.rgb = ACCENT2_COLOR

    doc.save(str(output_path))
    print(f"  Reference template: {output_path}")


def insert_cover_page(doc, cover_path):
    """Insert a full-page cover image as the first page with zero margins."""
    body = doc.element.body

    img = Image.open(cover_path)
    img_w, img_h = img.size

    page_w_emu = THEME["page"]["w"] * EMU_PER_TWIP
    page_h_emu = THEME["page"]["h"] * EMU_PER_TWIP

    img_ratio = img_w / img_h
    page_ratio = THEME["page"]["w"] / THEME["page"]["h"]

    if img_ratio > page_ratio:
        display_h = page_h_emu
        display_w = int(display_h * img_ratio)
    else:
        display_w = page_w_emu
        display_h = int(display_w / img_ratio)

    # Clamp to page
    display_w = min(display_w, page_w_emu)
    display_h = min(display_h, page_h_emu)

    # Create cover paragraph at end (temp), add image, then move to front
    cover_para = doc.add_paragraph()
    cover_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    # Remove any spacing
    cover_para.paragraph_format.space_before = Pt(0)
    cover_para.paragraph_format.space_after = Pt(0)

    run = cover_para.add_run()
    run.add_picture(cover_path, width=Emu(display_w), height=Emu(display_h))

    # Move to beginning of body
    cover_elem = cover_para._element
    body.remove(cover_elem)
    body.insert(0, cover_elem)

    # Add section break (next page) with zero margins to isolate cover
    pPr = cover_elem.find(qn('w:pPr'))
    if pPr is None:
        pPr = OxmlElement('w:pPr')
        cover_elem.insert(0, pPr)

    sectPr = OxmlElement('w:sectPr')

    sect_type = OxmlElement('w:type')
    sect_type.set(qn('w:val'), 'nextPage')
    sectPr.append(sect_type)

    pgSz = OxmlElement('w:pgSz')
    pgSz.set(qn('w:w'), str(THEME["page"]["w"]))
    pgSz.set(qn('w:h'), str(THEME["page"]["h"]))
    sectPr.append(pgSz)

    pgMar = OxmlElement('w:pgMar')
    for attr in ['top', 'bottom', 'left', 'right', 'header', 'footer', 'gutter']:
        pgMar.set(qn('w:{}'.format(attr)), '0')
    sectPr.append(pgMar)

    pPr.append(sectPr)


def strip_cover_line(content):
    """Remove cover image markdown line from content."""
    lines = content.split('\n')
    return '\n'.join(
        l for l in lines
        if not re.match(r'\s*!\[.*\]\(.*cover.*\)', l, re.IGNORECASE)
    )


def convert_to_docx(md_file, docx_output, project_dir, ref_docx, cover_path):
    """Convert markdown to styled DOCX with cover page."""
    # Strip cover from markdown
    with open(md_file, 'r', encoding='utf-8') as f:
        content = f.read()
    stripped = strip_cover_line(content)

    with tempfile.NamedTemporaryFile(mode='w', suffix='.md', delete=False, encoding='utf-8') as f:
        f.write(stripped)
        tmp_md = f.name

    try:
        cmd = [
            'pandoc', tmp_md,
            '-f', 'markdown-yaml_metadata_block',
            '-o', str(docx_output),
            f'--reference-doc={ref_docx}',
            '--resource-path', str(project_dir),
            '--wrap=none',
        ]
        result = subprocess.run(cmd, capture_output=True, text=True,
                                cwd=str(project_dir), timeout=600)
        if result.returncode != 0:
            stderr = result.stderr.strip()
            if stderr:
                # Only show first few warnings
                lines = stderr.split('\n')
                shown = lines[:3]
                if len(lines) > 3:
                    shown.append(f"  ... and {len(lines)-3} more warnings")
                print('\n'.join(f"    {l}" for l in shown))
    finally:
        os.unlink(tmp_md)

    # Post-process: insert cover page
    if cover_path and cover_path.exists():
        doc = Document(str(docx_output))
        insert_cover_page(doc, str(cover_path))
        doc.save(str(docx_output))


def convert_to_pdf(md_file, pdf_output, project_dir, cover_path):
    """Convert markdown to double-spaced PDF with cover using pandoc+weasyprint."""
    with open(md_file, 'r', encoding='utf-8') as f:
        content = f.read()

    stripped = strip_cover_line(content)

    # Prepend cover image as full-page HTML
    cover_html = ""
    if cover_path and cover_path.exists():
        abs_cover = cover_path.resolve()
        cover_html = f'<img class="cover-image" src="file://{abs_cover}" alt="Cover">\n\n'

    final_md = cover_html + stripped

    with tempfile.NamedTemporaryFile(mode='w', suffix='.md', delete=False, encoding='utf-8') as f:
        f.write(final_md)
        tmp_md = f.name

    with tempfile.NamedTemporaryFile(mode='w', suffix='.css', delete=False, encoding='utf-8') as f:
        f.write(PDF_CSS)
        tmp_css = f.name

    try:
        cmd = [
            'pandoc', tmp_md,
            '-f', 'markdown-yaml_metadata_block',
            '-o', str(pdf_output),
            '--pdf-engine=weasyprint',
            f'--css={tmp_css}',
            '--resource-path', str(project_dir),
            '--wrap=none',
        ]
        result = subprocess.run(cmd, capture_output=True, text=True,
                                cwd=str(project_dir), timeout=600)
        if result.returncode != 0:
            stderr = result.stderr.strip()
            # Filter out CSS warnings from weasyprint
            errors = [l for l in stderr.split('\n') if 'ERROR' in l.upper() or 'fatal' in l.lower()]
            if errors:
                for e in errors[:3]:
                    print(f"    {e}")
                return False
        return True
    finally:
        os.unlink(tmp_md)
        os.unlink(tmp_css)


def get_all_documents():
    """Discover all compilable project documents."""
    docs = []
    for project_dir in sorted(PROJECTS_DIR.iterdir()):
        if not project_dir.is_dir():
            continue
        output_dir = project_dir / "output"
        if not output_dir.exists():
            continue

        md_files = sorted(output_dir.glob("*.md"))
        if not md_files:
            continue

        cover = project_dir / "content" / "art" / "cover.png"
        has_cover = cover.exists()

        for md_file in md_files:
            docs.append({
                'project': project_dir.name,
                'project_dir': project_dir,
                'md_file': md_file,
                'stem': md_file.stem,
                'cover': cover if has_cover else None,
            })

    return docs


def main():
    print("=" * 70)
    print("  Supplement Converter")
    print("  Theme-Driven Styling + Triple-Spaced PDF")
    print("=" * 70)

    # Create reference template (ephemeral; regenerated from the active theme
    # on every run, not checked into the repo)
    with tempfile.TemporaryDirectory() as ref_tmpdir:
        ref_docx = Path(ref_tmpdir) / "reference.docx"
        print("\nCreating theme-driven reference template...")
        create_reference_docx(ref_docx)

        # Discover documents
        docs = get_all_documents()
        project_names = sorted(set(d['project'] for d in docs))
        print(f"\nFound {len(docs)} documents across {len(project_names)} projects:")
        for name in project_names:
            count = sum(1 for d in docs if d['project'] == name)
            has_cover = any(d['cover'] for d in docs if d['project'] == name)
            cover_status = "cover" if has_cover else "NO COVER"
            print(f"  {name} ({count} doc{'s' if count>1 else ''}, {cover_status})")

        # Process each document
        success_docx = 0
        success_pdf = 0
        failures = []

        for i, doc_info in enumerate(docs, 1):
            project = doc_info['project']
            stem = doc_info['stem']
            md_file = doc_info['md_file']
            project_dir = doc_info['project_dir']
            cover = doc_info['cover']
            output_dir = project_dir / "output"

            docx_out = output_dir / f"{stem}.docx"
            pdf_out = output_dir / f"{stem}.pdf"

            print(f"\n[{i}/{len(docs)}] {project}/{stem}")

            # DOCX
            try:
                print(f"  DOCX: converting...")
                convert_to_docx(md_file, docx_out, project_dir, ref_docx, cover)
                size_mb = docx_out.stat().st_size / (1024 * 1024)
                print(f"  DOCX: {docx_out.name} ({size_mb:.1f} MB)")
                success_docx += 1
            except Exception as e:
                print(f"  DOCX FAILED: {e}")
                failures.append((project, stem, 'docx', str(e)))

            # PDF
            try:
                print(f"  PDF:  converting (triple-spaced)...")
                ok = convert_to_pdf(md_file, pdf_out, project_dir, cover)
                if ok and pdf_out.exists():
                    size_mb = pdf_out.stat().st_size / (1024 * 1024)
                    print(f"  PDF:  {pdf_out.name} ({size_mb:.1f} MB)")
                    success_pdf += 1
                else:
                    print(f"  PDF:  conversion returned errors")
                    failures.append((project, stem, 'pdf', 'conversion failed'))
            except Exception as e:
                print(f"  PDF FAILED: {e}")
                failures.append((project, stem, 'pdf', str(e)))

    # Summary
    print(f"\n{'=' * 70}")
    print(f"  COMPLETE: {success_docx}/{len(docs)} DOCX, {success_pdf}/{len(docs)} PDF")
    if failures:
        print(f"\n  Failures:")
        for proj, stem, fmt, err in failures:
            print(f"    {proj}/{stem} ({fmt}): {err}")
    print(f"{'=' * 70}")


if __name__ == '__main__':
    main()
